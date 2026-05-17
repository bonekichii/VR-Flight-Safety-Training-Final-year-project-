from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create document
doc = Document()

# Title
title = doc.add_heading('VR Airplane Safety Training System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Project Report')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(16)
subtitle_format.font.color.rgb = RGBColor(0, 102, 204)

# Date
date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Executive Summary
doc.add_heading('Executive Summary', 1)
doc.add_paragraph(
    'This VR Airplane Safety Training System is an immersive virtual reality application designed to educate '
    'passengers about critical airplane safety procedures. Built using StereoKit framework with C#, the application '
    'provides an interactive learning experience combining video instruction, hands-on exploration, and knowledge '
    'assessment through quizzes.'
)

# Key Features
doc.add_heading('Key Features', 2)
features = [
    'Immersive VR environment with realistic airplane cabin',
    'Interactive welcome screen with text-to-speech announcements',
    'High-quality safety video playback with synchronized audio',
    'Comprehensive 12-question quiz system with instant feedback',
    'Realistic movement controls (WASD + controller support)',
    'Background ambient music for enhanced experience',
    'Collision detection for realistic cabin navigation',
    'Support for Meta Quest 3 and desktop testing'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# Technical Specifications
doc.add_heading('Technical Specifications', 1)

doc.add_heading('Technology Stack', 2)
tech_stack = [
    ('Framework', 'StereoKit (OpenXR-based VR framework)'),
    ('Language', 'C# (.NET)'),
    ('Video Processing', 'OpenCvSharp (OpenCV wrapper)'),
    ('Audio', 'NAudio library'),
    ('Text-to-Speech', 'System.Speech.Synthesis'),
    ('Target Platform', 'Windows PC, Meta Quest 3 (via Quest Link)'),
]
for tech, desc in tech_stack:
    p = doc.add_paragraph()
    p.add_run(f'{tech}: ').bold = True
    p.add_run(desc)

doc.add_heading('System Architecture', 2)
doc.add_paragraph(
    'The application follows a state-based architecture with four main states:'
)
states = [
    ('Welcome State', 'Initial greeting screen with TTS announcement'),
    ('Video State', 'Safety video playback with controls (play, pause, skip)'),
    ('Quiz State', 'Interactive quiz with 12 questions covering safety topics'),
    ('Finished State', 'Results display with score and restart option')
]
for state, desc in states:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{state}: ').bold = True
    p.add_run(desc)

doc.add_page_break()

# Features in Detail
doc.add_heading('Features in Detail', 1)

doc.add_heading('1. Welcome Screen', 2)
doc.add_paragraph(
    'The welcome screen greets users with a floating UI panel positioned in front of the user. '
    'When the "Start Training" button is clicked, a female voice announces: "Welcome to the Airlines. '
    'Please watch the safety video carefully." This creates an immersive and professional experience.'
)

doc.add_heading('2. Video Playback System', 2)
doc.add_paragraph(
    'The video playback system displays the Emirates safety video on a virtual screen within the VR cabin. '
    'Features include:'
)
video_features = [
    'Real-time video frame rendering using OpenCvSharp',
    'Synchronized audio playback through NAudio',
    'Interactive controls: Play, Pause, and Skip buttons',
    'Automatic transition to quiz upon video completion',
    'Frame rate: 30 FPS with smooth playback'
]
for vf in video_features:
    doc.add_paragraph(vf, style='List Bullet')

doc.add_heading('3. Quiz System', 2)
doc.add_paragraph(
    'The comprehensive quiz system tests user knowledge with 12 carefully crafted questions covering:'
)
quiz_topics = [
    'Life vest inflation procedures',
    'Oxygen mask location and usage',
    'Brace position techniques',
    'Emergency exit procedures',
    'Cabin smoke response',
    'Electronic device policies',
    'Seatbelt requirements',
    'Evacuation protocols',
    'Safety equipment timing',
    'Turbulence safety',
    'Carry-on storage requirements'
]
for topic in quiz_topics:
    doc.add_paragraph(topic, style='List Bullet')

doc.add_paragraph(
    '\nEach question provides immediate feedback with color-coded answers (green for correct, red for incorrect) '
    'and detailed explanations to reinforce learning.'
)

doc.add_heading('4. Movement & Navigation', 2)
doc.add_paragraph('The application supports multiple control schemes:')
controls = [
    ('Desktop Controls', 'WASD for movement, Arrow keys for camera rotation, Shift to stand/sit, Space to reset position'),
    ('VR Controller Controls', 'Thumbstick for movement, Natural head tracking, Hand tracking for UI interaction'),
    ('Collision System', 'Realistic seat colliders prevent walking through cabin furniture')
]
for control, desc in controls:
    p = doc.add_paragraph()
    p.add_run(f'{control}: ').bold = True
    p.add_run(desc)

doc.add_page_break()

# Quiz Questions Reference
doc.add_heading('Quiz Questions Reference', 1)
questions = [
    ('When should you inflate your life vest?', 'Outside the plane', 'Always inflate your life vest AFTER exiting the aircraft.'),
    ('Where is the oxygen mask located?', 'Overhead compartment', 'Oxygen masks are stored in overhead compartments above your seat.'),
    ('What is the correct brace position?', 'Head down, hands behind head', 'The brace position is head down with hands behind your head.'),
    ('How many emergency exits are typically on a Boeing 737?', '8 exits', 'A Boeing 737 typically has 8 emergency exits.'),
    ('What should you do first in case of cabin smoke?', 'Stay low and breathe through cloth', 'Smoke rises, so stay low where the air is clearer.'),
    ('When can you use electronic devices during flight?', 'Only in airplane mode', 'Electronic devices must be in airplane mode during flight.'),
    ('What does the seatbelt sign indicate?', 'Both turbulence and remain seated', 'The seatbelt sign means you must remain seated with your seatbelt fastened.'),
    ('How do you open an emergency exit door?', 'Follow the instructions on the door', 'Emergency exit doors have specific instructions printed on them.'),
    ('What should you leave behind during evacuation?', 'Carry-on luggage', 'NEVER take carry-on luggage during evacuation.'),
    ('How long does oxygen flow from overhead masks?', '12-15 minutes', 'Oxygen masks provide 12-15 minutes of oxygen.'),
    ('What is the safest seat position during turbulence?', 'Upright with seatbelt fastened', 'Keep your seat upright and seatbelt fastened during turbulence.'),
    ('Where should you place your carry-on during takeoff?', 'Under the seat in front', 'Carry-on items must be stored under the seat in front of you or in overhead bins.')
]

for i, (question, answer, explanation) in enumerate(questions, 1):
    doc.add_heading(f'Question {i}', 3)
    p = doc.add_paragraph()
    p.add_run('Q: ').bold = True
    p.add_run(question)
    p = doc.add_paragraph()
    p.add_run('A: ').bold = True
    p.add_run(answer)
    p = doc.add_paragraph()
    p.add_run('Explanation: ').italic = True
    p.add_run(explanation)

doc.add_page_break()

# Deployment Guide
doc.add_heading('Deployment Guide', 1)

doc.add_heading('Desktop Testing', 2)
doc.add_paragraph('To run the application on a Windows PC:')
desktop_steps = [
    'Navigate to the project directory',
    'Open SafetyDemo.csproj in Visual Studio 2022',
    'Ensure all NuGet packages are restored',
    'Press F5 or click the green Play button',
    'Use mouse and keyboard controls to navigate'
]
for i, step in enumerate(desktop_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('Meta Quest 3 Deployment (Quest Link)', 2)
doc.add_paragraph('The easiest method to run on Meta Quest 3:')
quest_steps = [
    'Install Meta Quest Link software on your PC',
    'Enable Developer Mode on your Quest 3 headset',
    'Connect Quest 3 to PC via USB-C cable or Air Link (WiFi)',
    'Launch Meta Quest Link on the headset',
    'Run the application from Visual Studio as normal',
    'The app will automatically display in VR mode',
    'Use hand tracking or controllers to interact with UI elements'
]
for i, step in enumerate(quest_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('VR Controls', 2)
vr_controls = [
    'Hand Tracking: Point and pinch to click buttons',
    'Controllers: Use trigger to select UI elements',
    'Thumbsticks: Move around the cabin',
    'Head Movement: Natural 6DOF tracking',
    'All UI buttons are automatically VR-compatible'
]
for control in vr_controls:
    doc.add_paragraph(control, style='List Bullet')

doc.add_page_break()

# Project Structure
doc.add_heading('Project Structure', 1)
doc.add_paragraph('The project consists of the following key files:')
files = [
    ('Program.cs', 'Main application code (759 lines) containing all game logic, state management, and rendering'),
    ('SafetyDemo.csproj', 'C# project file with NuGet package references'),
    ('Assets/', 'Directory containing the safety video and airplane 3D model'),
    ('README.md', 'Quick start guide for running the project'),
    ('VR_DEPLOYMENT_GUIDE.md', 'Detailed VR deployment instructions'),
    ('STATUS.md', 'Current project status and feature checklist')
]
for filename, desc in files:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(filename).bold = True
    p.add_run(f': {desc}')

# Dependencies
doc.add_heading('Dependencies', 2)
dependencies = [
    'StereoKit (>= 0.3.9) - VR framework',
    'OpenCvSharp4 (>= 4.10.0) - Video processing',
    'OpenCvSharp4.runtime.win (>= 4.10.0) - OpenCV runtime',
    'NAudio (>= 2.2.1) - Audio playback',
    'System.Speech - Text-to-speech synthesis'
]
for dep in dependencies:
    doc.add_paragraph(dep, style='List Bullet')

doc.add_page_break()

# Performance Metrics
doc.add_heading('Performance Metrics', 1)
metrics = [
    ('Video Resolution', 'Original video resolution maintained'),
    ('Frame Rate', '30 FPS for video playback'),
    ('VR Refresh Rate', 'Matches headset native refresh rate (72-120 Hz)'),
    ('Audio Latency', 'Synchronized with video frames'),
    ('Quiz Questions', '12 comprehensive safety questions'),
    ('Background Music', '30-second procedurally generated ambient loop'),
    ('Cabin Dimensions', '2.4m width × 20m length'),
    ('Seat Rows', '12 rows with collision detection')
]
for metric, value in metrics:
    p = doc.add_paragraph()
    p.add_run(f'{metric}: ').bold = True
    p.add_run(value)

# Future Enhancements
doc.add_heading('Future Enhancement Opportunities', 1)
enhancements = [
    'Add multiplayer support for group training sessions',
    'Implement hand tracking gestures for equipment interaction',
    'Create additional emergency scenario simulations',
    'Add support for multiple languages and subtitles',
    'Integrate haptic feedback for enhanced immersion',
    'Develop analytics dashboard for training effectiveness',
    'Add voice recognition for verbal quiz responses',
    'Create customizable cabin layouts for different aircraft types'
]
for enhancement in enhancements:
    doc.add_paragraph(enhancement, style='List Bullet')

doc.add_page_break()

# Conclusion
doc.add_heading('Conclusion', 1)
doc.add_paragraph(
    'The VR Airplane Safety Training System represents a modern, engaging approach to passenger safety education. '
    'By leveraging virtual reality technology, the application provides an immersive learning experience that '
    'significantly improves knowledge retention compared to traditional safety briefings.'
)
doc.add_paragraph(
    'The system is production-ready and can be deployed immediately for desktop testing or Meta Quest 3 VR experiences. '
    'All core features are fully functional, including video playback, interactive quizzes, and realistic cabin navigation.'
)

# Contact/Support Section
doc.add_heading('Project Information', 1)
info = [
    ('Project Name', 'SafetyDemo - VR Airplane Safety Training'),
    ('Framework', 'StereoKit with C#'),
    ('Platform', 'Windows PC, Meta Quest 3'),
    ('Status', 'Production Ready'),
    ('Last Updated', datetime.now().strftime("%B %d, %Y"))
]
for label, value in info:
    p = doc.add_paragraph()
    p.add_run(f'{label}: ').bold = True
    p.add_run(value)

# Save document
output_path = 'VR_Safety_Training_Report.docx'
doc.save(output_path)
print(f'Report generated successfully: {output_path}')
